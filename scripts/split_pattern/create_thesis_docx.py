#!/usr/bin/env python3
"""
生成示例论文正文.docx — 匿名 scaffold
格式严格参照苏州大学本科毕业论文规范。
"""

import os
import sys

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import lxml.etree as etree
import re

# ============================================================
# 路径常量
# ============================================================
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_PATH = os.path.join(SCRIPT_DIR, '示例论文正文.docx')
FIG_DIR = os.path.join(SCRIPT_DIR, 'figures')

# 从 content_data.py 导入全部论文内容
sys.path.insert(0, SCRIPT_DIR)
from content_data import (
    TITLE_CN, TITLE_EN, ABSTRACT_CN, KEYWORDS_CN, ABSTRACT_EN, KEYWORDS_EN,
    CH1_SECTIONS, CH2_SECTIONS,
    CH3_SEC_31, CH3_SEC_32, CH3_SEC_33, CH3_SEC_34, CH3_SEC_35, CH3_SEC_36, CH3_SEC_37,
    CH4_SECTIONS, CH5_PARAS,
    REFERENCES, ACKNOWLEDGEMENT,
    FIGURE_MAP, FIGURE_CAPTIONS, TOC_ENTRIES,
)

# 页面设置
MARGIN_TOP = Cm(3.3)
MARGIN_BOTTOM = Cm(2.7)
MARGIN_LEFT = Cm(2.5)
MARGIN_RIGHT = Cm(2.5)
GUTTER = Cm(0.5)
HEADER_DISTANCE = Cm(2.6)
FOOTER_DISTANCE = Cm(2.0)

# ============================================================
# 文档格式函数（复用文献综述格式）
# ============================================================

def setup_page(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = MARGIN_TOP
    section.bottom_margin = MARGIN_BOTTOM
    section.left_margin = MARGIN_LEFT
    section.right_margin = MARGIN_RIGHT
    section.header_distance = HEADER_DISTANCE
    section.footer_distance = FOOTER_DISTANCE
    sectPr = section._sectPr
    pgMar = sectPr.find(qn('w:pgMar'))
    if pgMar is not None:
        pgMar.set(qn('w:gutter'), str(int(GUTTER.emu / 635)))


def setup_doc_defaults(doc):
    styles_element = doc.styles.element
    docDefaults = styles_element.find(qn('w:docDefaults'))
    if docDefaults is None:
        docDefaults = parse_xml(
            f'<w:docDefaults {nsdecls("w")}>'
            f'<w:rPrDefault><w:rPr>'
            f'<w:rFonts w:ascii="Calibri" w:eastAsia="宋体" w:hAnsi="Calibri" w:cs="Times New Roman"/>'
            f'<w:sz w:val="21"/><w:szCs w:val="22"/>'
            f'</w:rPr></w:rPrDefault>'
            f'</w:docDefaults>'
        )
        styles_element.insert(0, docDefaults)
    else:
        rPrDefault = docDefaults.find(qn('w:rPrDefault'))
        if rPrDefault is None:
            rPrDefault = parse_xml(
                f'<w:rPrDefault {nsdecls("w")}><w:rPr>'
                f'<w:rFonts w:ascii="Calibri" w:eastAsia="宋体" w:hAnsi="Calibri" w:cs="Times New Roman"/>'
                f'<w:sz w:val="21"/><w:szCs w:val="22"/>'
                f'</w:rPr></w:rPrDefault>'
            )
            docDefaults.insert(0, rPrDefault)


def setup_styles(doc):
    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = 'Times New Roman'
    normal.font.size = Pt(12)
    normal_el = normal.element
    rpr = normal_el.find(qn('w:rPr'))
    if rpr is None:
        rpr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        normal_el.append(rpr)
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rpr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), '宋体')
    pf = normal.paragraph_format
    pf.line_spacing = 1.5
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    # Heading 1: 15pt(小三), 粗体, 居中
    h1 = styles['Heading 1']
    h1.font.name = 'Times New Roman'
    h1.font.size = Pt(15)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.line_spacing = 1.5
    h1_pPr = h1.element.find(qn('w:pPr'))
    if h1_pPr is None:
        h1_pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
        h1.element.insert(0, h1_pPr)
    h1_spacing = h1_pPr.find(qn('w:spacing'))
    if h1_spacing is None:
        h1_spacing = parse_xml(f'<w:spacing {nsdecls("w")}/>')
        h1_pPr.append(h1_spacing)
    h1_spacing.set(qn('w:beforeLines'), '100')
    h1_spacing.set(qn('w:afterLines'), '100')
    for attr in ['w:before', 'w:after']:
        if h1_spacing.get(qn(attr)) is not None:
            del h1_spacing.attrib[qn(attr)]
    h1_rpr = h1.element.find(qn('w:rPr'))
    if h1_rpr is None:
        h1_rpr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        h1.element.append(h1_rpr)
    h1_rf = h1_rpr.find(qn('w:rFonts'))
    if h1_rf is None:
        h1_rf = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        h1_rpr.append(h1_rf)
    h1_rf.set(qn('w:eastAsia'), '宋体')

    # Heading 2: 12pt(小四), 粗体, 左对齐
    h2 = styles['Heading 2']
    h2.font.name = 'Times New Roman'
    h2.font.size = Pt(12)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2.paragraph_format.line_spacing = 1.5
    h2_pPr = h2.element.find(qn('w:pPr'))
    if h2_pPr is None:
        h2_pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
        h2.element.insert(0, h2_pPr)
    h2_spacing = h2_pPr.find(qn('w:spacing'))
    if h2_spacing is None:
        h2_spacing = parse_xml(f'<w:spacing {nsdecls("w")}/>')
        h2_pPr.append(h2_spacing)
    h2_spacing.set(qn('w:beforeLines'), '100')
    h2_spacing.set(qn('w:afterLines'), '100')
    for attr in ['w:before', 'w:after']:
        if h2_spacing.get(qn(attr)) is not None:
            del h2_spacing.attrib[qn(attr)]
    h2_rpr = h2.element.find(qn('w:rPr'))
    if h2_rpr is None:
        h2_rpr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        h2.element.append(h2_rpr)
    h2_rf = h2_rpr.find(qn('w:rFonts'))
    if h2_rf is None:
        h2_rf = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        h2_rpr.append(h2_rf)
    h2_rf.set(qn('w:eastAsia'), '宋体')

    # ---------- TOC 样式（完全匹配参考论文） ----------
    from docx.enum.style import WD_STYLE_TYPE
    # 直接操作 styles.xml 创建内置TOC样式（不使用python-docx的add_style以避免customStyle属性）
    styles_element = doc.styles.element

    def _ensure_toc_style(styles_element, style_id, style_xml):
        """确保TOC样式存在，如果已存在则替换，否则新增"""
        existing = styles_element.find(f'{{{qn("w:style").split("}")[0][1:]}}}style[@{qn("w:styleId")}="{style_id}"]')
        if existing is None:
            for s in styles_element.findall(qn('w:style')):
                if s.get(qn('w:styleId')) == style_id:
                    existing = s
                    break
        new_elem = parse_xml(style_xml)
        if existing is not None:
            existing.getparent().replace(existing, new_elem)
        else:
            styles_element.append(new_elem)

    # TOC 1: 四号(14pt=sz28), Times New Roman/黑体, bCs, 段前段后120twip
    _ensure_toc_style(styles_element, 'TOC1', (
        f'<w:style {nsdecls("w")} w:type="paragraph" w:styleId="TOC1">'
        f'<w:name w:val="toc 1"/>'
        f'<w:basedOn w:val="Normal"/>'
        f'<w:next w:val="Normal"/>'
        f'<w:autoRedefine/>'
        f'<w:uiPriority w:val="39"/>'
        f'<w:unhideWhenUsed/>'
        f'<w:qFormat/>'
        f'<w:pPr>'
        f'<w:tabs><w:tab w:val="right" w:leader="dot" w:pos="8776"/></w:tabs>'
        f'<w:spacing w:before="120" w:after="120"/>'
        f'</w:pPr>'
        f'<w:rPr>'
        f'<w:rFonts w:ascii="Times New Roman" w:eastAsia="\u9ed1\u4f53" w:hAnsi="Times New Roman"/>'
        f'<w:bCs/>'
        f'<w:sz w:val="28"/><w:szCs w:val="20"/>'
        f'</w:rPr>'
        f'</w:style>'
    ))

    # TOC 2: 四号(14pt=sz28), Times New Roman/宋体, 左缩进240, 行距360auto
    _ensure_toc_style(styles_element, 'TOC2', (
        f'<w:style {nsdecls("w")} w:type="paragraph" w:styleId="TOC2">'
        f'<w:name w:val="toc 2"/>'
        f'<w:basedOn w:val="Normal"/>'
        f'<w:next w:val="Normal"/>'
        f'<w:autoRedefine/>'
        f'<w:uiPriority w:val="39"/>'
        f'<w:unhideWhenUsed/>'
        f'<w:pPr>'
        f'<w:tabs><w:tab w:val="right" w:leader="dot" w:pos="8776"/></w:tabs>'
        f'<w:spacing w:line="360" w:lineRule="auto"/>'
        f'<w:ind w:left="240"/>'
        f'</w:pPr>'
        f'<w:rPr>'
        f'<w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>'
        f'<w:sz w:val="28"/><w:szCs w:val="20"/>'
        f'</w:rPr>'
        f'</w:style>'
    ))


def set_run_font_eastasia(run, font_name):
    run_el = run._element
    rpr = run_el.find(qn('w:rPr'))
    if rpr is None:
        rpr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        run_el.insert(0, rpr)
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rpr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)


def add_header_footer(doc):
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = hp.add_run("苏州大学本科生毕业设计（论文）")
    run.font.size = Pt(9)
    run.font.name = '宋体'
    set_run_font_eastasia(run, '宋体')

    # 页眉底部横线
    hp_pPr = hp._element.find(qn('w:pPr'))
    if hp_pPr is None:
        hp_pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
        hp._element.insert(0, hp_pPr)
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="auto"/>'
        f'</w:pBdr>'
    )
    hp_pPr.append(pBdr)

    # 页脚: -N- 格式页码
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add "-" before page number
    run_dash1 = fp.add_run("-")
    run_dash1.font.size = Pt(9)
    run_dash1.font.name = 'Times New Roman'

    # PAGE field
    run_f1 = fp.add_run()
    run_f1._element.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))
    run_f2 = fp.add_run()
    run_f2._element.append(parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'))
    run_f3 = fp.add_run()
    run_f3._element.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))

    # Add "-" after page number
    run_dash2 = fp.add_run("-")
    run_dash2.font.size = Pt(9)
    run_dash2.font.name = 'Times New Roman'


def add_heading(doc, text, level, add_toc_bookmark=True, page_break_before=False):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'Times New Roman'
        set_run_font_eastasia(run, '宋体')
    if page_break_before:
        # 用 pageBreakBefore 属性而非空段落+分页符，避免产生空白页
        pPr = heading._element.find(qn('w:pPr'))
        if pPr is None:
            pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
            heading._element.insert(0, pPr)
        pPr.append(parse_xml(f'<w:pageBreakBefore {nsdecls("w")}/>'))
    if add_toc_bookmark:
        _add_toc_bookmark_to_heading(heading)
    return heading


def add_non_chapter_heading(doc, text, add_toc_bookmark=True, page_break_before=False):
    """添加非章节标题（参考文献、致谢等），外观与Heading 1相同但不使用Heading样式，
    不会被Word编入章节大纲。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    # 段前/段后间距使用"行"单位（beforeLines/afterLines），而非固定磅值
    pPr = p._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
        p._element.insert(0, pPr)
    nch_sp = pPr.find(qn('w:spacing'))
    if nch_sp is None:
        nch_sp = parse_xml(f'<w:spacing {nsdecls("w")}/>')
        pPr.append(nch_sp)
    nch_sp.set(qn('w:beforeLines'), '100')
    nch_sp.set(qn('w:afterLines'), '100')

    if page_break_before:
        pPr.append(parse_xml(f'<w:pageBreakBefore {nsdecls("w")}/>'))

    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(15)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0, 0, 0)
    set_run_font_eastasia(run, '宋体')

    if add_toc_bookmark:
        _add_toc_bookmark_to_heading(p)

    return p


CITATION_RE = re.compile(r'(\[\d+(?:[-,]\d+)*\])')


def _add_run_with_font(p, text, superscript=False):
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    set_run_font_eastasia(run, '宋体')
    if superscript:
        run.font.superscript = True
    return run


def _add_citation_hyperlink(p, citation_text):
    num_match = re.search(r'\d+', citation_text)
    if not num_match:
        _add_run_with_font(p, citation_text, superscript=True)
        return
    ref_num = num_match.group()
    anchor_name = f'_Ref{ref_num}'
    nsmap = 'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
    rpr_xml = (
        f'<w:rPr {nsmap}>'
        f'<w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:eastAsia="宋体"/>'
        f'<w:color w:val="000000"/>'
        f'<w:sz w:val="24"/><w:szCs w:val="24"/>'
        f'<w:u w:val="none"/>'
        f'<w:vertAlign w:val="superscript"/>'
        f'</w:rPr>'
    )
    hyperlink_xml = (
        f'<w:hyperlink {nsmap} w:anchor="{anchor_name}">'
        f'<w:r>{rpr_xml}<w:t>{citation_text}</w:t></w:r>'
        f'</w:hyperlink>'
    )
    hyperlink_elem = etree.fromstring(hyperlink_xml.encode('utf-8'))
    p._element.append(hyperlink_elem)


def add_body_paragraph(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pPr = p._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
        p._element.insert(0, pPr)
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = parse_xml(f'<w:ind {nsdecls("w")}/>')
        pPr.append(ind)
    ind.set(qn('w:firstLineChars'), '200')

    parts = CITATION_RE.split(text)
    for part in parts:
        if not part:
            continue
        if CITATION_RE.fullmatch(part):
            _add_citation_hyperlink(p, part)
        else:
            _add_run_with_font(p, part)
    return p


def add_figure(doc, fig_path, caption_text, width=Inches(5.0)):
    if not os.path.exists(fig_path):
        print(f"  [警告] 图片未找到: {fig_path}")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"[图片待插入: {os.path.basename(fig_path)}]")
        run.font.color.rgb = RGBColor(255, 0, 0)
    else:
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.line_spacing = 1.0
        p_img.paragraph_format.space_before = Pt(6)
        run = p_img.add_run()
        run.add_picture(fig_path, width=width)

    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p_cap.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.5
    run_cap = p_cap.add_run(caption_text)
    run_cap.font.size = Pt(12)
    run_cap.font.name = 'Times New Roman'
    set_run_font_eastasia(run_cap, '宋体')


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    br = parse_xml(f'<w:br {nsdecls("w")} w:type="page"/>')
    run._element.append(br)


# 全局TOC书签计数器
_toc_bookmark_counter = 0


def _next_toc_id():
    global _toc_bookmark_counter
    _toc_bookmark_counter += 1
    return _toc_bookmark_counter


def _build_toc_entry_xml(text, anchor, page_num, is_level1):
    """构建一条TOC条目的完整XML段落（使用pStyle引用TOC样式，匹配参考论文）"""
    ns = 'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
    style_id = 'TOC1' if is_level1 else 'TOC2'
    # 文本run用 noProof; 页码run用 noProof + webHidden（匹配参考论文）
    rpr_text = f'<w:rPr {ns}><w:noProof/></w:rPr>'
    rpr_page = f'<w:rPr {ns}><w:noProof/><w:webHidden/></w:rPr>'
    safe_text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
    # toc 1 需要段落级行距360（匹配参考论文）
    if is_level1:
        ppr = f'<w:pPr><w:pStyle w:val="{style_id}"/><w:spacing w:line="360" w:lineRule="auto"/></w:pPr>'
    else:
        ppr = f'<w:pPr><w:pStyle w:val="{style_id}"/></w:pPr>'
    xml = (
        f'<w:p {ns}>'
        f'{ppr}'
        f'<w:hyperlink w:anchor="{anchor}">'
        f'<w:r>{rpr_text}<w:t xml:space="preserve">{safe_text}</w:t></w:r>'
        f'<w:r>{rpr_page}<w:tab/></w:r>'
        f'<w:r>{rpr_page}<w:fldChar w:fldCharType="begin"/></w:r>'
        f'<w:r>{rpr_page}<w:instrText xml:space="preserve"> PAGEREF {anchor} \\h </w:instrText></w:r>'
        f'<w:r>{rpr_page}<w:fldChar w:fldCharType="separate"/></w:r>'
        f'<w:r>{rpr_page}<w:t>{page_num}</w:t></w:r>'
        f'<w:r>{rpr_page}<w:fldChar w:fldCharType="end"/></w:r>'
        f'</w:hyperlink>'
        f'</w:p>'
    )
    return xml


def _replace_paragraph_xml(doc, new_xml_bytes):
    """用doc.add_paragraph()创建段落，然后替换其XML为自定义内容"""
    placeholder = doc.add_paragraph()
    new_elem = parse_xml(new_xml_bytes)
    placeholder._element.getparent().replace(placeholder._element, new_elem)
    return new_elem


def add_toc(doc):
    """插入目录：TOC域代码 + hyperlink条目 + PAGEREF（匹配参考论文格式）"""
    global _toc_bookmark_counter
    _toc_bookmark_counter = 0
    ns = 'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'

    toc_entries = TOC_ENTRIES

    for idx, (text, page, is_level1) in enumerate(toc_entries):
        toc_id = _next_toc_id()
        anchor = f'_Toc{toc_id}'
        entry_xml = _build_toc_entry_xml(text, anchor, page, is_level1)
        entry_elem = _replace_paragraph_xml(doc, entry_xml.encode('utf-8'))

        if idx == 0:
            # 把TOC域开始插入到第一个条目段落的最前面（不创建额外段落）
            fld_sep = parse_xml(f'<w:r {ns}><w:fldChar w:fldCharType="separate"/></w:r>')
            fld_instr = parse_xml(f'<w:r {ns}><w:instrText xml:space="preserve"> TOC \\o "1-2" \\h \\z \\u </w:instrText></w:r>')
            fld_begin = parse_xml(f'<w:r {ns}><w:fldChar w:fldCharType="begin"/></w:r>')
            # 插入顺序：begin, instrText, separate（在段落pPr之后、hyperlink之前）
            first_child = entry_elem[0]  # pPr
            first_child.addnext(fld_sep)
            first_child.addnext(fld_instr)
            first_child.addnext(fld_begin)

        if idx == len(toc_entries) - 1:
            # 把TOC域结束追加到最后一个条目段落的末尾
            fld_end = parse_xml(f'<w:r {ns}><w:fldChar w:fldCharType="end"/></w:r>')
            entry_elem.append(fld_end)


def _add_toc_bookmark_to_heading(heading_para):
    """给标题段落添加对应的_Toc书签"""
    toc_id = _next_toc_id()
    bm_name = f'_Toc{toc_id}'
    bm_id_num = 200 + toc_id
    add_bookmark_to_paragraph(heading_para, bm_name, bm_id_num)


def add_bookmark_to_paragraph(para, bookmark_name, bookmark_id):
    p_elem = para._element
    runs = p_elem.findall(qn('w:r'))
    if runs:
        bm_start = parse_xml(
            f'<w:bookmarkStart {nsdecls("w")} w:id="{bookmark_id}" w:name="{bookmark_name}"/>'
        )
        runs[0].addprevious(bm_start)
        bm_end = parse_xml(
            f'<w:bookmarkEnd {nsdecls("w")} w:id="{bookmark_id}"/>'
        )
        runs[-1].addnext(bm_end)


def add_references(doc, refs):
    for ref in refs:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing = 1.5
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pPr = p._element.find(qn('w:pPr'))
        if pPr is None:
            pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
            p._element.insert(0, pPr)
        ind = pPr.find(qn('w:ind'))
        if ind is None:
            ind = parse_xml(f'<w:ind {nsdecls("w")}/>')
            pPr.append(ind)
        ind.set(qn('w:hanging'), '720')
        ind.set(qn('w:left'), '720')

        ref_text = re.sub(r'^\[(\d+)\] ', r'[\1]\t', ref)
        run = p.add_run(ref_text)
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        set_run_font_eastasia(run, '宋体')

        ref_match = re.match(r'\[(\d+)\]', ref)
        if ref_match:
            ref_num = ref_match.group(1)
            bm_name = f'_Ref{ref_num}'
            bm_id = 100 + int(ref_num)
            add_bookmark_to_paragraph(p, bm_name, bm_id)


# ============================================================
# ============================================================
# 论文内容从 content_data.py 导入
# ============================================================

# 主文档生成函数
# ============================================================

def create_document():
    print("=" * 60)
    print("  生成示例论文正文.docx")
    print("  " + TITLE_CN)
    print("=" * 60)

    doc = Document()

    # 1. 页面设置
    print("[1/12] 页面设置...")
    setup_page(doc)
    setup_doc_defaults(doc)
    setup_styles(doc)
    add_header_footer(doc)

    # ========== 中文摘要 ==========
    print("[2/12] 中文摘要...")
    # 标题
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p_title.paragraph_format
    pf.space_before = Pt(17)
    pf.space_after = Pt(16.5)
    pf.line_spacing = 1.5
    run = p_title.add_run(TITLE_CN)
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = 'Times New Roman'
    set_run_font_eastasia(run, '黑体')

    # "摘 要"标题 — 用普通段落+手动格式化，避免Heading样式自带的小黑点
    p_abs = doc.add_paragraph()
    p_abs.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_abs.paragraph_format.line_spacing = 1.5
    # 段前/段后间距使用"行"单位
    _pPr_abs = p_abs._element.find(qn('w:pPr'))
    if _pPr_abs is None:
        _pPr_abs = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
        p_abs._element.insert(0, _pPr_abs)
    _sp_abs = _pPr_abs.find(qn('w:spacing'))
    if _sp_abs is None:
        _sp_abs = parse_xml(f'<w:spacing {nsdecls("w")}/>')
        _pPr_abs.append(_sp_abs)
    _sp_abs.set(qn('w:beforeLines'), '100')
    _sp_abs.set(qn('w:afterLines'), '100')
    run_abs = p_abs.add_run("摘 要")
    run_abs.bold = True
    run_abs.font.size = Pt(15)
    run_abs.font.name = 'Times New Roman'
    run_abs.font.color.rgb = RGBColor(0, 0, 0)
    set_run_font_eastasia(run_abs, '宋体')

    # 摘要正文 - 按段落拆分
    for para_text in ABSTRACT_CN.split('\n'):
        para_text = para_text.strip()
        if para_text:
            add_body_paragraph(doc, para_text)

    # 关键词
    p_kw = doc.add_paragraph()
    p_kw.paragraph_format.line_spacing = 1.5
    pPr = p_kw._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
        p_kw._element.insert(0, pPr)
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = parse_xml(f'<w:ind {nsdecls("w")}/>')
        pPr.append(ind)
    ind.set(qn('w:firstLineChars'), '200')
    run_kl = p_kw.add_run("关键词：")
    run_kl.bold = True
    run_kl.font.size = Pt(12)
    run_kl.font.name = 'Times New Roman'
    set_run_font_eastasia(run_kl, '宋体')
    run_kv = p_kw.add_run(KEYWORDS_CN)
    run_kv.font.size = Pt(12)
    run_kv.font.name = 'Times New Roman'
    set_run_font_eastasia(run_kv, '宋体')

    # ========== English Abstract ==========
    print("[3/12] English Abstract...")
    add_page_break(doc)

    # 英文标题
    p_title_en = doc.add_paragraph()
    p_title_en.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf_te = p_title_en.paragraph_format
    pf_te.space_before = Pt(17)
    pf_te.space_after = Pt(16.5)
    pf_te.line_spacing = 1.5
    run_te = p_title_en.add_run(TITLE_EN)
    run_te.bold = True
    run_te.font.size = Pt(18)
    run_te.font.name = 'Times New Roman'
    set_run_font_eastasia(run_te, '黑体')

    # Abstract标题 — 同样避免Heading样式的小黑点
    p_abs_en = doc.add_paragraph()
    p_abs_en.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_abs_en.paragraph_format.line_spacing = 1.5
    # 段前/段后间距使用"行"单位
    _pPr_abs_en = p_abs_en._element.find(qn('w:pPr'))
    if _pPr_abs_en is None:
        _pPr_abs_en = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
        p_abs_en._element.insert(0, _pPr_abs_en)
    _sp_abs_en = _pPr_abs_en.find(qn('w:spacing'))
    if _sp_abs_en is None:
        _sp_abs_en = parse_xml(f'<w:spacing {nsdecls("w")}/>')
        _pPr_abs_en.append(_sp_abs_en)
    _sp_abs_en.set(qn('w:beforeLines'), '100')
    _sp_abs_en.set(qn('w:afterLines'), '100')
    run_abs_en = p_abs_en.add_run("Abstract")
    run_abs_en.bold = True
    run_abs_en.font.size = Pt(15)
    run_abs_en.font.name = 'Times New Roman'
    run_abs_en.font.color.rgb = RGBColor(0, 0, 0)
    set_run_font_eastasia(run_abs_en, '宋体')

    for para_text in ABSTRACT_EN.split('\n'):
        para_text = para_text.strip()
        if para_text:
            add_body_paragraph(doc, para_text)

    p_kw_en = doc.add_paragraph()
    p_kw_en.paragraph_format.line_spacing = 1.5
    pPr2 = p_kw_en._element.find(qn('w:pPr'))
    if pPr2 is None:
        pPr2 = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
        p_kw_en._element.insert(0, pPr2)
    ind2 = pPr2.find(qn('w:ind'))
    if ind2 is None:
        ind2 = parse_xml(f'<w:ind {nsdecls("w")}/>')
        pPr2.append(ind2)
    ind2.set(qn('w:firstLineChars'), '200')
    run_kl2 = p_kw_en.add_run("Keywords: ")
    run_kl2.bold = True
    run_kl2.font.size = Pt(12)
    run_kl2.font.name = 'Times New Roman'
    set_run_font_eastasia(run_kl2, '宋体')
    run_kv2 = p_kw_en.add_run(KEYWORDS_EN)
    run_kv2.font.size = Pt(12)
    run_kv2.font.name = 'Times New Roman'
    set_run_font_eastasia(run_kv2, '宋体')

    # ========== 目录 ==========
    print("[4/13] 目录...")
    add_page_break(doc)
    # "目 录"标题 — 用普通段落，不进入TOC本身
    p_toc_title = doc.add_paragraph()
    p_toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_toc_title.paragraph_format.line_spacing = 1.5
    # 段前/段后间距使用"行"单位
    _pPr_toc = p_toc_title._element.find(qn('w:pPr'))
    if _pPr_toc is None:
        _pPr_toc = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
        p_toc_title._element.insert(0, _pPr_toc)
    _sp_toc = _pPr_toc.find(qn('w:spacing'))
    if _sp_toc is None:
        _sp_toc = parse_xml(f'<w:spacing {nsdecls("w")}/>')
        _pPr_toc.append(_sp_toc)
    _sp_toc.set(qn('w:beforeLines'), '100')
    _sp_toc.set(qn('w:afterLines'), '100')
    run_toc = p_toc_title.add_run("目 录")
    run_toc.bold = True
    run_toc.font.size = Pt(14)
    run_toc.font.name = 'Times New Roman'
    run_toc.font.color.rgb = RGBColor(0, 0, 0)
    set_run_font_eastasia(run_toc, '宋体')
    add_toc(doc)

    # 重置TOC书签计数器，使正文标题书签ID与目录条目ID匹配
    global _toc_bookmark_counter
    _toc_bookmark_counter = 0

    # ========== 第1章 绪论 ==========
    print("[5/13] 第1章 绪论...")
    add_heading(doc, "第1章 绪论", level=1, page_break_before=True)
    for section_title, paragraphs in CH1_SECTIONS.items():
        add_heading(doc, section_title, level=2)
        for para in paragraphs:
            add_body_paragraph(doc, para)

    # ========== 第2章 材料与方法 ==========
    print("[6/13] 第2章 材料与方法...")
    add_heading(doc, "第2章 材料与方法", level=1, page_break_before=True)
    for section_title, paragraphs in CH2_SECTIONS.items():
        add_heading(doc, section_title, level=2)
        for para in paragraphs:
            add_body_paragraph(doc, para)

    # ========== 第3章 结果 ==========
    print("[7/13] 第3章 结果...")
    add_heading(doc, "第3章 结果", level=1, page_break_before=True)

    # 3.1 示例资料整理结果
    add_heading(doc, "3.1 示例资料整理结果", level=2)
    for para in CH3_SEC_31:
        add_body_paragraph(doc, para)
    add_figure(doc, FIGURE_MAP["fig3_1"], FIGURE_CAPTIONS["fig3_1"])
    add_figure(doc, FIGURE_MAP["fig3_2"], FIGURE_CAPTIONS["fig3_2"])

    # 3.2 示例描述性统计结果
    add_heading(doc, "3.2 示例描述性统计结果", level=2)
    for para in CH3_SEC_32:
        add_body_paragraph(doc, para)
    add_figure(doc, FIGURE_MAP["fig3_3"], FIGURE_CAPTIONS["fig3_3"])
    add_figure(doc, FIGURE_MAP["fig3_4"], FIGURE_CAPTIONS["fig3_4"])

    # 3.3 示例分组分析结果
    add_heading(doc, "3.3 示例分组分析结果", level=2)
    for para in CH3_SEC_33:
        add_body_paragraph(doc, para)
    add_figure(doc, FIGURE_MAP["fig3_5"], FIGURE_CAPTIONS["fig3_5"])
    add_figure(doc, FIGURE_MAP["fig3_6"], FIGURE_CAPTIONS["fig3_6"])

    # 3.4 示例子指标分析结果
    add_heading(doc, "3.4 示例子指标分析结果", level=2)
    for para in CH3_SEC_34:
        add_body_paragraph(doc, para)
    add_figure(doc, FIGURE_MAP["fig3_7"], FIGURE_CAPTIONS["fig3_7"])
    add_figure(doc, FIGURE_MAP["fig3_8"], FIGURE_CAPTIONS["fig3_8"])

    # 3.5 示例综合指标分析结果
    add_heading(doc, "3.5 示例综合指标分析结果", level=2)
    for para in CH3_SEC_35:
        add_body_paragraph(doc, para)
    add_figure(doc, FIGURE_MAP["fig3_9"], FIGURE_CAPTIONS["fig3_9"])
    add_figure(doc, FIGURE_MAP["fig3_10"], FIGURE_CAPTIONS["fig3_10"])

    # 3.6 示例模型验证结果
    add_heading(doc, "3.6 示例模型验证结果", level=2)
    for para in CH3_SEC_36:
        add_body_paragraph(doc, para)
    add_figure(doc, FIGURE_MAP["fig3_11"], FIGURE_CAPTIONS["fig3_11"])
    add_figure(doc, FIGURE_MAP["fig3_12"], FIGURE_CAPTIONS["fig3_12"])

    # 3.7 示例交叉验证结果
    add_heading(doc, "3.7 示例交叉验证结果", level=2)
    for para in CH3_SEC_37:
        add_body_paragraph(doc, para)
    add_figure(doc, FIGURE_MAP["fig3_13"], FIGURE_CAPTIONS["fig3_13"])

    # ========== 第4章 讨论 ==========
    print("[8/13] 第4章 讨论...")
    add_heading(doc, "第4章 讨论", level=1, page_break_before=True)
    for section_title, paragraphs in CH4_SECTIONS.items():
        add_heading(doc, section_title, level=2)
        for para in paragraphs:
            add_body_paragraph(doc, para)

    # ========== 第5章 结论 ==========
    print("[9/13] 第5章 结论...")
    add_heading(doc, "第5章 结论", level=1, page_break_before=True)
    for para in CH5_PARAS:
        add_body_paragraph(doc, para)

    # ========== 参考文献 ==========
    print("[10/13] 参考文献...")
    add_non_chapter_heading(doc, "参考文献", page_break_before=True)
    add_references(doc, REFERENCES)

    # ========== 致谢 ==========
    print("[11/13] 致谢...")
    add_non_chapter_heading(doc, "致 谢", page_break_before=True)
    for para_text in ACKNOWLEDGEMENT.split('\n'):
        para_text = para_text.strip()
        if para_text:
            add_body_paragraph(doc, para_text)

    # ========== 清理元数据（去除python-docx痕迹） ==========
    print("[12/13] 文档设置...")
    import datetime
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    core = doc.core_properties
    core.title = ''
    core.subject = ''
    core.author = ''
    core.last_modified_by = ''
    core.keywords = ''
    core.category = ''
    core.comments = ''
    core.content_status = ''
    core.identifier = ''
    core.language = ''
    core.version = ''
    core.revision = 1
    core.created = datetime.datetime.now()
    core.modified = datetime.datetime.now()
    # 替换/清空 app.xml 元数据，避免留下生成工具、公司或管理者信息
    app_part = doc.part.package.part_related_by(RT.EXTENDED_PROPERTIES)
    if app_part is not None:
        from lxml import etree as _et
        app_root = _et.fromstring(app_part.blob)
        ns = {'ep': 'http://schemas.openxmlformats.org/officeDocument/2006/extended-properties'}
        for tag_name in ['Application', 'AppVersion', 'Company', 'Manager']:
            el = app_root.find(f'ep:{tag_name}', ns)
            if el is not None:
                if tag_name == 'Application':
                    el.text = 'Microsoft Office Word'
                elif tag_name == 'AppVersion':
                    el.text = '16.0000'
                else:
                    el.text = ''
        app_part._blob = _et.tostring(app_root, xml_declaration=True, encoding='UTF-8', standalone=True)

    # ========== 保存 ==========
    print("[13/13] 保存文档...")
    doc.save(OUTPUT_PATH)
    file_size = os.path.getsize(OUTPUT_PATH)

    # 统计字数
    total_chars = len(ABSTRACT_CN) + len(KEYWORDS_CN) + len(ABSTRACT_EN) + len(KEYWORDS_EN)
    total_chars += len(ACKNOWLEDGEMENT)
    for sections in [CH1_SECTIONS, CH2_SECTIONS, CH4_SECTIONS]:
        for title, paras in sections.items():
            total_chars += len(title)
            for p in paras:
                total_chars += len(p)
    for sec_paras in [CH3_SEC_31, CH3_SEC_32, CH3_SEC_33, CH3_SEC_34, CH3_SEC_35,
                      CH3_SEC_36, CH3_SEC_37, CH5_PARAS]:
        for p in sec_paras:
            total_chars += len(p)
    for ref in REFERENCES:
        total_chars += len(ref)

    print(f"\n{'='*60}")
    print(f"  文档已保存: {OUTPUT_PATH}")
    print(f"  文件大小: {file_size / 1024:.1f} KB")
    print(f"  正文+引用总字符数: ~{total_chars}")
    print(f"  参考文献数: {len(REFERENCES)}条")
    print(f"{'='*60}")

    return OUTPUT_PATH


if __name__ == '__main__':
    create_document()
